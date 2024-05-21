import oracledb as ora
import json
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# SQL to fetch tablespace and data file metrics
TABLESPACE_INFO_SQL = """
    SELECT
    ts.tablespace_name,
    ts.block_size,
    ts.initial_extent,
    ts.next_extent,
    ts.min_extents,
    ts.max_extents,
    ts.pct_increase,
    ts.min_extlen,
    ts.status,
    ts.contents,
    ts.logging,
    ts.force_logging,
    ts.extent_management,
    ts.segment_space_management,
    df.file_name,
    df.bytes / 1024 / 1024 AS file_size_mb,
    NVL(fs.bytes, 0) / 1024 / 1024 AS free_space_mb,
    (df.bytes - NVL(fs.bytes, 0)) / 1024 / 1024 AS used_space_mb,
    df.autoextensible,
    df.maxbytes / 1024 / 1024 AS max_file_size_mb
    FROM
        dba_tablespaces ts
        LEFT JOIN dba_data_files df ON ts.tablespace_name = df.tablespace_name
        LEFT JOIN (
            SELECT
                tablespace_name,
                file_id,
                SUM(bytes) AS bytes
            FROM
                dba_free_space
            GROUP BY
                tablespace_name,
                file_id
        ) fs ON df.file_id = fs.file_id
    ORDER BY
        ts.tablespace_name,
        df.file_name
"""

# SQL to calculate fragmentation
FRAGMENTATION_SQL = """
    SELECT
    tablespace_name,
    COUNT(*) AS fragments,
    MAX(bytes) / 1024 / 1024 AS max_fragment_size_mb
    FROM
    dba_free_space
    GROUP BY
    tablespace_name
    ORDER BY
    tablespace_name
"""

# SQL to fetch number of users
USERS_SQL = """
    SELECT COUNT(*) AS number_of_users FROM dba_users
"""

# SQL to fetch schema sizes
SCHEMA_SIZE_SQL = """
    SELECT
    owner AS schema_name,
    SUM(bytes) / 1024 / 1024 AS schema_size_mb
    FROM
    dba_segments
    GROUP BY
    owner
    ORDER BY
    owner
"""
# SQL to fetch ASM disk group usage
ASM_DISKGROUP_USAGE_SQL = """
    SELECT
        g.name AS diskgroup_name,
        g.total_mb,
        g.free_mb,
        g.total_mb - g.free_mb AS used_mb,
        ROUND((1 - g.free_mb / g.total_mb) * 100, 2) AS used_pct
    FROM
        v$asm_diskgroup g
"""

def get_asm_diskgroup_usage(conn):
    try:
        cur = conn.cursor()
        cur.execute(ASM_DISKGROUP_USAGE_SQL)
        rows = cur.fetchall()
        headers = [column[0] for column in cur.description]
        return [dict(zip(headers, row)) for row in rows]
    except Exception as e:
        print(f"Error fetching ASM disk group usage: {e}")
        return None


def get_database_info(database):
    try:
        with ora.connect(
            service_name=database["name"],
            host=database["host"],
            port=database["port"],
            user=database["username"],
            password=database["password"],
            mode=ora.SYSDBA
        ) as conn:
            cur = conn.cursor()
            
            # Fetch tablespace information
            cur.execute(TABLESPACE_INFO_SQL)
            tablespace_rows = cur.fetchall()
            tablespace_headers = [column[0] for column in cur.description]
            tablespace_info = [dict(zip(tablespace_headers, row)) for row in tablespace_rows]

            # Fetch fragmentation information
            cur.execute(FRAGMENTATION_SQL)
            fragmentation_rows = cur.fetchall()
            fragmentation_headers = [column[0] for column in cur.description]
            fragmentation_info = [dict(zip(fragmentation_headers, row)) for row in fragmentation_rows]
            asm_diskgroup_usage = get_asm_diskgroup_usage(conn)

            # Merge fragmentation info into tablespace info
            for ts_info in tablespace_info:
                ts_name = ts_info['TABLESPACE_NAME']
                frag_info = next((frag for frag in fragmentation_info if frag['TABLESPACE_NAME'] == ts_name), None)
                if frag_info:
                    ts_info['FRAGMENTS'] = frag_info['FRAGMENTS']
                    ts_info['MAX_FRAGMENT_SIZE_MB'] = frag_info['MAX_FRAGMENT_SIZE_MB']
                else:
                    ts_info['FRAGMENTS'] = 0
                    ts_info['MAX_FRAGMENT_SIZE_MB'] = 0

            # Fetch number of users
            cur.execute(USERS_SQL)
            number_of_users = cur.fetchone()[0]

            # Fetch schema sizes
            cur.execute(SCHEMA_SIZE_SQL)
            schema_rows = cur.fetchall()
            schema_headers = [column[0] for column in cur.description]
            schema_info = [dict(zip(schema_headers, row)) for row in schema_rows]

            return {
                "tablespace_info": tablespace_info,
                "number_of_users": number_of_users,
                "schema_info": schema_info,
                "asm_diskgroup_usage": asm_diskgroup_usage
            }
    except Exception as e:
        print(f"Error fetching data from database {database['name']}: {e}")
        return None

def get_all_databases_info(databases):
    all_databases_info = []
    for database in databases:
        database_info = get_database_info(database)
        if database_info:
            all_databases_info.append({
                "database_name": database["name"],
                "info": database_info
            })
    return all_databases_info

def generate_pie_chart(df: pd.DataFrame, column, labels, title, legend_title, img_path, threshold):
    # Calculate the 8% threshold
    threshold = df[column].sum() * (threshold/100)

    # Create a new DataFrame for "Others"
    others_df = df[df[column] < threshold].copy()
    others_df.loc['Others', column] = others_df[column].sum()
    others_df.loc['Others', labels.name] = 'Others'

    # Remove the rows that are now in "Others"
    df = df[df[column] >= threshold]

    # Add the "Others" row to the DataFrame
    df = pd.concat([df, pd.DataFrame(others_df.loc['Others']).T])

    fig, ax = plt.subplots(figsize=(10, 6))
    wedges, texts, autotexts = ax.pie(df[column], labels=df[labels.name], autopct='%1.1f%%', startangle=90)
    ax.axis('equal')
    plt.title(title)
    plt.legend(wedges, df[labels.name], title=legend_title, loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
    plt.savefig(img_path)
    plt.close(fig)

def add_tablespace_details_table_to_doc(doc, df: pd.DataFrame):
    # Transpose the DataFrame and reset the index
    df = df.transpose().reset_index()

    # Replace NaN values with a default value
    df.fillna('N/A', inplace=True)

    # Rename the columns
    df.columns = df.iloc[0]
    df = df[1:]

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)  # Convert column names to strings

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)  # Convert values to strings

def add_schema_details_table_to_doc(doc, df: pd.DataFrame):
    # Sort the DataFrame by schema size
    df.sort_values(by='SCHEMA_SIZE_MB', ascending=False, inplace=True)

    # Replace NaN values with a default value
    df.fillna('N/A', inplace=True)

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)  # Convert column names to strings

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)  # Convert values to strings


def add_table_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)


def generate_reports(databases_info):
    doc = Document()
    doc.add_heading('Oracle Database Monthly Utilization Report', 0)
    threshold = 5

    for database_info in databases_info:
        db_name = database_info["database_name"]
        info = database_info["info"]
        tablespace_info = info["tablespace_info"]
        number_of_users = info["number_of_users"]
        schema_info = info["schema_info"]
        asm_diskgroup_usage = info["asm_diskgroup_usage"]

        doc.add_heading(f'Database: {db_name}', level=1)

        # Tablespace Information
        doc.add_heading('Tablespace Information', level=2)
        df = pd.DataFrame(tablespace_info)
        df['USED_SPACE_MB'].fillna(0, inplace=True)
        df['FREE_SPACE_MB'].fillna(0, inplace=True)
        df['PERCENTAGE_USED'] = df['USED_SPACE_MB'] / (df['USED_SPACE_MB'] + df['FREE_SPACE_MB']) * 100

        img_path = f'used_space_{db_name}.png'
        generate_pie_chart(df, 'USED_SPACE_MB', df['TABLESPACE_NAME'], f'Used Space by Tablespace for {db_name}', 'Tablespaces', img_path, threshold)
        doc.add_picture(img_path, width=Inches(6))

        # Tablespace details table
        doc.add_heading('Tablespace Details', level=3)
        add_tablespace_details_table_to_doc(doc, df)  # Transpose the tablespace details

        # Number of Users
        doc.add_heading('Number of Users', level=2)
        doc.add_paragraph(f'Total number of users: {number_of_users}')

        # Schema Sizes
        doc.add_heading('Schema Sizes', level=2)
        schema_df = pd.DataFrame(schema_info)

        img_path = f'schema_sizes_{db_name}.png'
        generate_pie_chart(schema_df, 'SCHEMA_SIZE_MB', schema_df['SCHEMA_NAME'], f'Schema Sizes for {db_name}', 'Schemas', img_path,threshold)
        doc.add_picture(img_path, width=Inches(6))

        # Schema details table
        add_schema_details_table_to_doc(doc, schema_df)  # Do not transpose the schema details
        # ASM Disk Group Usage
        doc.add_heading('ASM Disk Group Usage', level=2)
        asm_df = pd.DataFrame(asm_diskgroup_usage)
        add_table_to_doc(doc, asm_df)

    report_path = 'oracle_utilization_report.docx'
    doc.save(report_path)
    print(f'Report generated: {report_path}')


def main():
    with open("oracle.json") as conf_file:
        conf = json.load(conf_file)
        databases = conf["databases"]
        databases_info = get_all_databases_info(databases)
        generate_reports(databases_info)

if __name__ == "__main__":
    main()
